
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

import pandas as pd

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
from pipeline import config

OUT = Path(__file__).resolve().parent
MD = OUT / "MY_RESEARCH_PAPER.md"
DOCX = OUT / "MY_RESEARCH_PAPER.docx"

ACCENT = RGBColor(0x1A, 0x3D, 0x6D)
INK = RGBColor(0x1A, 0x1A, 0x19)
MUTED = RGBColor(0x55, 0x55, 0x50)

BOLD = re.compile(r"\*\*(.+?)\*\*")
ITAL = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
CODE = re.compile(r"`([^`]+?)`")
IMG = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")


def add_rich(par, text):
    tokens = re.split(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)", text)
    for t in tokens:
        if not t:
            continue
        if BOLD.fullmatch(t):
            r = par.add_run(BOLD.fullmatch(t).group(1)); r.bold = True
        elif CODE.fullmatch(t):
            r = par.add_run(CODE.fullmatch(t).group(1))
            r.font.name = "Consolas"; r.font.size = Pt(8.5)
        elif ITAL.fullmatch(t):
            r = par.add_run(ITAL.fullmatch(t).group(1)); r.italic = True
        else:
            par.add_run(t)


def parse_table(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(set(c) <= set("-: ") for c in cells if c):
            rows.append(cells)
        i += 1
    return rows, i


def main():
    lines = MD.read_text(encoding="utf-8").split("\n")
    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)
    st.paragraph_format.space_after = Pt(6)
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(0.9)
        sec.top_margin = sec.bottom_margin = Inches(0.85)

    i = 0
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()

        if not s:
            i += 1
            continue

        m = IMG.match(s)
        if m:
            p = OUT / m.group(2)
            if p.exists():
                doc.add_picture(str(p), width=Inches(6.4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if s.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            par = doc.add_paragraph()
            r = par.add_run("\n".join(buf))
            r.font.name = "Consolas"; r.font.size = Pt(8.5)
            par.paragraph_format.left_indent = Inches(0.25)
            continue

        if s.startswith("|"):
            rows, i = parse_table(lines, i)
            if not rows:
                continue
            ncol = max(len(r) for r in rows)
            t = doc.add_table(rows=0, cols=ncol)
            t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                cells = t.add_row().cells
                for ci in range(ncol):
                    txt = row[ci] if ci < len(row) else ""
                    cp = cells[ci].paragraphs[0]
                    cp.paragraph_format.space_after = Pt(2)
                    add_rich(cp, txt)
                    for run in cp.runs:
                        run.font.size = Pt(8)
                        if ri == 0:
                            run.bold = True
            doc.add_paragraph()
            continue

        if s.startswith("#"):
            lvl = len(s) - len(s.lstrip("#"))
            txt = s.lstrip("#").strip()
            if lvl == 1:
                h = doc.add_heading(level=0)
                r = h.add_run(txt); r.font.color.rgb = ACCENT; r.font.size = Pt(17)
            else:
                h = doc.add_heading(level=min(lvl - 1, 4))
                r = h.add_run(txt)
                r.font.color.rgb = ACCENT
                r.font.size = Pt({2: 13, 3: 11, 4: 10}.get(lvl, 10))
            i += 1
            continue

        if s.startswith("- ") or re.match(r"^\d+\.\s", s):
            style = "List Bullet" if s.startswith("- ") else "List Number"
            body = s[2:] if s.startswith("- ") else re.sub(r"^\d+\.\s", "", s)
            par = doc.add_paragraph(style=style)
            add_rich(par, body)
            i += 1
            continue

        if s.startswith("---"):
            i += 1
            continue

        par = doc.add_paragraph()
        add_rich(par, s.rstrip("  "))
        if s.startswith("*") and s.endswith("*"):
            for r in par.runs:
                r.font.color.rgb = MUTED
                r.font.size = Pt(8.5)
        i += 1

    doc.save(DOCX)
    print(f"  wrote {DOCX.name}")

    AV = json.loads((OUT / "audit_verification.json").read_text(encoding="utf-8"))
    E76 = json.loads((config.EXPERIMENTS_DIR /
                      "exp_76_architectural_diversity_blend.json").read_text(encoding="utf-8"))
    E77 = json.loads((config.EXPERIMENTS_DIR /
                      "exp_77_recursive_member_upgrade.json").read_text(encoding="utf-8"))
    E79 = json.loads((config.EXPERIMENTS_DIR /
                      "exp_79_upgrade_seed_check.json").read_text(encoding="utf-8"))
    BL = AV["SHIPPED blend w=0.60"]
    OP = pd.DataFrame(E77["operating_point"])
    OP = OP[OP.pair == "AB2"]

    S = []
    a = S.append
    a("# Independent Technical Assessment — NPN_HACKATHON")
    a("")
    a("Companion to `MY_RESEARCH_PAPER.pdf`. Every figure below was "
      "re-derived from project artifacts during the audit; none is taken on trust.")
    a("")
    a("## Ratings")
    a("")
    a("| Dimension | Rating |")
    a("|---|---|")
    a("| Overall model quality | **Moderate** |")
    a("| Validation quality | **Good** |")
    a("| Leakage status | **Clean (independently verified)** |")
    a("| Reproducibility | **Good** |")
    a("")
    a("*Moderate*, not *Good*, on model quality: the gains are real and "
      "replicated, but small in absolute terms on a noisy target, and the "
      "shipped configuration regresses MAE.")
    a("")
    a("## Final model")
    a("")
    a("```")
    a("y_hat = 0.60 * Direct(38 features)  +  0.40 * Recursive(32 features)")
    a("")
    a("  Direct    : LightGBM Tweedie(p=1.1), 400 rounds, 15 origins x 28 days")
    a("  Recursive : LightGBM Tweedie(p=1.1), 400 rounds, 420 daily origins,")
    a("              rolled forward 28 days on its own output")
    a("  Weight    : selected on inner window d_1886-d_1913 (pre-origin)")
    a("```")
    a("")
    a("## Verified performance — primary window, 853,720 predictions")
    a("")
    a("| Metric | Value |")
    a("|---|---|")
    a(f"| RMSE | **{BL['RMSE']:.4f}** |")
    a(f"| MAE | **{BL['MAE']:.4f}** |")
    a(f"| WAPE | {BL['WAPE']:.4f} |")
    a(f"| Bias | {BL['bias']:+.4f} |")
    a(f"| High-volume RMSE | {BL['high_volume_RMSE']:.4f} |")
    a(f"| Demand-occurrence accuracy (y>0, thr 0.5) | {BL['Accuracy']:.4f} |")
    a(f"| Precision / Recall / F1 | {BL['Precision']:.4f} / {BL['Recall']:.4f} / {BL['F1']:.4f} |")
    a("")
    a(f"Independently reproduced from raw CSVs: RMSE "
      f"{AV['reproduction']['measured_RMSE']:.4f} vs recorded "
      f"{AV['reproduction']['expected_RMSE']:.4f} "
      f"(drift {abs(AV['reproduction']['measured_RMSE']-AV['reproduction']['expected_RMSE']):.1e}).")
    a("")
    a(f"Across four disjoint windows: mean ΔRMSE **{OP.dRMSE_vs_A.mean():+.4f}**, "
      f"mean ΔMAE **{OP.dMAE_vs_A.mean():+.4f}** versus the direct member alone.")
    a("")
    a("## Leakage status")
    a("")
    a("| Test | Result |")
    a("|---|---|")
    a(f"| Future sales corrupted → features change? | **0 of "
      f"{AV['leakage_test']['features_checked']}** → PASS |")
    a("| Future prices corrupted → price features change? | Yes → PASS (mirror test) |")
    a("| Recursive rollout reads post-origin actuals? | No — structurally impossible, verified |")
    a("| Train/validation target overlap | None (train ends d_1913, validation starts d_1914) |")
    a("| Target encoding / global normalisation | Not used |")
    a("| Blend weight or hyperparameters fitted on evaluation data | No |")
    a("")
    a("**Conclusion: no target leakage.** Safety is structural — lags are "
      "origin-relative by construction — and empirically verified, not asserted.")
    a("")
    a("## Defects found")
    a("")
    a("1. **Mislabelled prediction file.** "
      "`predictions/validation/exp_74_new_champion_validation.csv` is byte-identical "
      "to `exp_72_shape_validation.csv` (36-feature model), not the 38-feature "
      "champion. Registry metrics unaffected; no downstream result depends on it.")
    a("2. **Registry does not name the shipped configuration.** Exp. #77's "
      "`metrics` field holds the w=0.50 acceptance-test blend (2.0915/1.0433), not "
      "the shipped w=0.60 (2.0929/1.0395, in `operating_point`).")
    a("3. **Seed-convention sensitivity.** Setting `bagging_seed`/"
      "`feature_fraction_seed` explicitly vs deriving them from `seed` shifts RMSE "
      "by up to 0.005 — larger than some accepted effects. Both conventions appear "
      "in the codebase.")
    a("4. **No predictions persisted for accepted models** (the shipped model had "
      "none until this audit regenerated them).")
    a("")
    a("None of these invalidates the reported performance.")
    a("")
    a("## Strongest contribution")
    a("")
    a("The negative-control attribution of ensemble gain: "
      f"{E76['negative_control']['same_architecture_gain']:+.4f} from averaging a "
      f"reseeded copy (residual corr "
      f"{E76['negative_control']['same_architecture_resid_corr']:.4f}) versus "
      f"{E76['negative_control']['diversity_gain']:+.4f} from a different "
      f"architecture (corr 0.9496) — "
      f"**{E76['negative_control']['gain_attributable_to_architecture']:+.4f} "
      "attributable to architecture**. This rescued a direction the project had "
      "already rejected, and diagnosed *why* the earlier attempt failed.")
    a("")
    a("## Biggest weakness")
    a("")
    a(f"The MAE regression ({OP.dMAE_vs_A.mean():+.4f}). The operating point was "
      "chosen to optimise RMSE without an explicit business loss function "
      "justifying that trade. Close behind: the final accepted gain is carried by "
      "two of four windows, and on one window the upgraded member was worse than "
      "the one it replaced.")
    a("")
    a("## What to improve next")
    a("")
    a("1. State the loss function; re-select the blend weight against it.")
    a("2. Persist predictions for every accepted model; fix the mislabelled file.")
    a("3. Standardise the seed convention across scripts.")
    a("4. If the business can re-forecast weekly, do that before any further "
      "modelling — it beats every remaining modelling idea.")
    a("")
    a("## Evidence trail")
    a("")
    a("| Artifact | Path |")
    a("|---|---|")
    a("| Audit reproduction + leakage test | `audit_reproduce.py`, `audit_verification.json` |")
    a("| Regenerated shipped predictions | `reproduction/shipped_blend_w060_validation.csv` |")
    a("| Verified metric table | `MODEL_COMPARISON.csv` |")
    a("| Figures + underlying tables | `figures/` |")
    a(f"| Experiment registry | `experiments/registry/` "
      f"({len(list(config.EXPERIMENTS_DIR.glob('*.json')))} records) |")
    a("")
    (OUT / "AUDIT_SUMMARY.md").write_text("\n".join(S), encoding="utf-8")
    print("  wrote AUDIT_SUMMARY.md")


if __name__ == "__main__":
    main()
